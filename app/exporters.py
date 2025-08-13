# app/exporters.py
from docx import Document
from docx.shared import Pt
from io import BytesIO

def build_docx(entries):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    doc.add_heading("Live Caption Transcript", level=1)
    for i, e in enumerate(entries, 1):
        doc.add_paragraph(f"[{i}] ({e['t0']:.2f}–{e['t1']:.2f}s)")
        doc.add_paragraph(f"EN: {e['text_en']}")
        if e["text_ko"]:
            doc.add_paragraph(f"KO: {e['text_ko']}")
        doc.add_paragraph("")  # spacing
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
